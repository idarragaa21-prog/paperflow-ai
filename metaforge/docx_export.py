"""Export a drafted manuscript to a Word (.docx) file, with embedded figures
(forest / PRISMA / risk-of-bias) and a GRADE certainty table."""
from __future__ import annotations

import io
import re

from .figures import svg_to_png
from .grade import DOMAIN_LABELS_ES, DOWNGRADE_DOMAINS
from .manuscript import SECTIONS

_SECTION_TITLES_ES = {
    "abstract": "Resumen", "introduction": "Introducción", "methods": "Métodos",
    "results": "Resultados", "discussion": "Discusión", "limitations": "Limitaciones",
    "conclusion": "Conclusión",
}
_DOWN_LABEL = {"none": "No", "serious": "Sí (−1)", "very_serious": "Muy serio (−2)"}


def _add_runs(paragraph, text: str) -> None:
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not seg:
            continue
        run = paragraph.add_run(seg)
        if i % 2 == 1:
            run.bold = True


def _add_markdown(doc, text: str) -> None:
    for raw in (text or "").split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.lstrip().startswith(("- ", "* ")):
            _add_runs(doc.add_paragraph(style="List Bullet"), line.lstrip()[2:].strip())
        else:
            _add_runs(doc.add_paragraph(), line)


def _add_figure(doc, svg: str | None, caption: str) -> bool:
    if not svg:
        return False
    png = svg_to_png(svg, width=1400)
    if not png:
        return False
    from docx.shared import Inches, Pt
    doc.add_picture(io.BytesIO(png), width=Inches(6.2))
    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    return True


def _add_grade_table(doc, grade: dict) -> None:
    from docx.shared import Pt
    cert = grade.get("certainty", {})
    doc.add_heading("Certeza de la evidencia (GRADE)", level=1)
    pico = doc.add_paragraph()
    r = pico.add_run(f"Certeza global: {cert.get('label', '—').upper()} (nivel {cert.get('level', '—')}/4).")
    r.bold = True
    domains = grade.get("domains")
    if not domains:
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Dominio", "Juicio (motivo)"
    for d in DOWNGRADE_DOMAINS:
        info = domains.get(d, {})
        cells = table.add_row().cells
        cells[0].text = DOMAIN_LABELS_ES.get(d, d)
        rating = _DOWN_LABEL.get(info.get("rating", "none"), info.get("rating", ""))
        cells[1].text = f"{rating} — {info.get('reason', '')}"
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)


def manuscript_docx(sections: dict, *, facts: str | None = None,
                    figures: dict | None = None, grade: dict | None = None) -> bytes:
    from docx import Document
    from docx.shared import Pt

    figures = figures or {}
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading((sections.get("title") or "Manuscrito").strip(), level=0)

    fig_no = [0]

    def next_fig():
        fig_no[0] += 1
        return fig_no[0]

    for sec in SECTIONS:
        if sec == "title":
            continue
        body = (sections.get(sec) or "").strip()
        if not body and sec not in ("methods", "results"):
            continue
        doc.add_heading(_SECTION_TITLES_ES.get(sec, sec.capitalize()), level=1)
        if body:
            _add_markdown(doc, body)
        if sec == "methods" and figures.get("prisma"):
            _add_figure(doc, figures["prisma"], f"Figura {next_fig()}. Diagrama de flujo PRISMA 2020.")
        if sec == "results":
            if figures.get("forest"):
                _add_figure(doc, figures["forest"], f"Figura {next_fig()}. Forest plot del efecto combinado.")
            if figures.get("funnel"):
                _add_figure(doc, figures["funnel"], f"Figura {next_fig()}. Funnel plot (sesgo de publicación).")
            if figures.get("rob"):
                _add_figure(doc, figures["rob"], f"Figura {next_fig()}. Evaluación del riesgo de sesgo.")

    if grade:
        _add_grade_table(doc, grade)
        if figures.get("sof"):
            _add_figure(doc, figures["sof"], "Tabla. Resumen de hallazgos (GRADE).")

    if facts:
        doc.add_page_break()
        doc.add_heading("Anexo: cifras verificadas", level=1)
        for line in facts.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
