"""Clinical sheet export services — DOCX and PDF.

DOCX: uses a Node.js script (scripts/generate_clinical_docx.js) that builds
      a professional, branded document with headers, footers, styled headings,
      inline bold/italic parsing, and a formatted references table.

PDF:  converts the generated DOCX to PDF via LibreOffice (soffice) when
      available; falls back to reportlab for plain-text PDF if soffice is absent.
"""
from __future__ import annotations

import json
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from uuid import UUID

from app.core.logger import logger
from app.core.storage import storage_manager

# Path to the Node script, relative to this file's package root.
_SCRIPT = Path(__file__).parents[3] / "scripts" / "generate_clinical_docx.js"


def _out_dir(user_id: UUID, sheet_id: UUID) -> Path:
    rel = Path("clinical_exports") / str(user_id) / str(sheet_id)
    abs_dir = (storage_manager.base_dir / rel).resolve()
    abs_dir.mkdir(parents=True, exist_ok=True)
    return abs_dir


def export_docx(
    *,
    user_id: UUID,
    sheet_id: UUID,
    version: int,
    topic: str,
    content_markdown: str,
    references_json: list | None = None,
    created_at: str | None = None,
    llm_model: str | None = None,
) -> str:
    """Generate a professional DOCX via Node/docx-js.

    Returns the relative path under storage_manager.base_dir.
    Falls back to a plain python-docx export if Node or the script is missing.
    """
    abs_dir = _out_dir(user_id, sheet_id)
    filename = f"clinical_{sheet_id}_v{int(version)}.docx"
    abs_path = abs_dir / filename

    # Prefer the professional generator when Node is available
    node_bin = shutil.which("node")
    if node_bin and _SCRIPT.exists():
        try:
            payload = {
                "topic": topic or "Clinical Sheet",
                "version": version,
                "created_at": created_at,
                "content_markdown": content_markdown or "",
                "references_json": references_json or [],
                "llm_model": llm_model or "",
            }
            with tempfile.NamedTemporaryFile(suffix=".json", delete=False, mode="w") as tmp:
                json.dump(payload, tmp, ensure_ascii=False)
                tmp_path = tmp.name

            result = subprocess.run(
                [node_bin, str(_SCRIPT), tmp_path, str(abs_path)],
                capture_output=True, text=True, timeout=60,
            )
            Path(tmp_path).unlink(missing_ok=True)

            if result.returncode == 0 and abs_path.exists():
                logger.info(f"[Exporters] DOCX generated via Node: {abs_path}")
                return str(abs_path.relative_to(storage_manager.base_dir))
            else:
                logger.warning(f"[Exporters] Node DOCX failed: {result.stderr[:400]}")
        except Exception as exc:
            logger.warning(f"[Exporters] Node DOCX exception: {exc}")

    # Fallback: python-docx (basic but always available)
    return _export_docx_fallback(
        user_id=user_id, sheet_id=sheet_id, version=version,
        topic=topic, content_markdown=content_markdown,
        abs_path=abs_path,
    )


def _export_docx_fallback(
    *,
    user_id: UUID,
    sheet_id: UUID,
    version: int,
    topic: str,
    content_markdown: str,
    abs_path: Path,
) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    text = re.sub(r"\n{3,}", "\n\n", (content_markdown or "").strip())
    doc = Document()

    # Cover title
    title = doc.add_heading(topic, level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # Meta line
    meta = doc.add_paragraph(f"PaperFlow AI · v{version}")
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    meta.runs[0].font.italic = True

    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("## "):
            doc.add_heading(block[3:], level=2)
        elif block.startswith("### "):
            doc.add_heading(block[4:], level=3)
        else:
            # Strip simple markdown and add paragraph
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", block)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            p = doc.add_paragraph(clean)
            p.runs[0].font.size = Pt(11) if p.runs else None

    doc.save(str(abs_path))
    logger.info(f"[Exporters] DOCX fallback saved: {abs_path}")
    return str(abs_path.relative_to(storage_manager.base_dir))


def export_pdf(
    *,
    user_id: UUID,
    sheet_id: UUID,
    version: int,
    topic: str,
    content_markdown: str,
    references_json: list | None = None,
    created_at: str | None = None,
    llm_model: str | None = None,
) -> str:
    """Export to PDF.

    Strategy:
      1. Generate DOCX (professional version if possible).
      2. Convert via LibreOffice soffice (best quality).
      3. Fall back to reportlab plain-text PDF.
    """
    abs_dir = _out_dir(user_id, sheet_id)
    pdf_filename = f"clinical_{sheet_id}_v{int(version)}.pdf"
    pdf_path = abs_dir / pdf_filename

    # Step 1: generate DOCX
    docx_rel = export_docx(
        user_id=user_id, sheet_id=sheet_id, version=version,
        topic=topic, content_markdown=content_markdown,
        references_json=references_json, created_at=created_at, llm_model=llm_model,
    )
    docx_abs = (storage_manager.base_dir / docx_rel).resolve()

    # Step 2: LibreOffice conversion
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice and docx_abs.exists():
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(abs_dir), str(docx_abs)],
                capture_output=True, text=True, timeout=120,
            )
            # soffice outputs with same stem
            converted = abs_dir / (docx_abs.stem + ".pdf")
            if result.returncode == 0 and converted.exists():
                if converted != pdf_path:
                    converted.rename(pdf_path)
                logger.info(f"[Exporters] PDF via LibreOffice: {pdf_path}")
                return str(pdf_path.relative_to(storage_manager.base_dir))
            else:
                logger.warning(f"[Exporters] LibreOffice failed: {result.stderr[:300]}")
        except Exception as exc:
            logger.warning(f"[Exporters] LibreOffice exception: {exc}")

    # Step 3: reportlab fallback
    return _export_pdf_reportlab(
        topic=topic, content_markdown=content_markdown,
        pdf_path=pdf_path, version=version,
    )


def _export_pdf_reportlab(*, topic: str, content_markdown: str, pdf_path: Path, version: int) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas
    from reportlab.lib import colors

    text = re.sub(r"\n{3,}", "\n\n", (content_markdown or "").strip())
    w, h = A4
    margin = 2 * cm
    c = canvas.Canvas(str(pdf_path), pagesize=A4)

    def new_page():
        c.showPage()
        # Header line
        c.setStrokeColorRGB(0.149, 0.227, 0.373)
        c.setLineWidth(1.5)
        c.line(margin, h - margin + 6, w - margin, h - margin + 6)

    y = h - margin - 20

    # Header
    c.setStrokeColorRGB(0.149, 0.227, 0.373)
    c.setLineWidth(1.5)
    c.line(margin, y + 10, w - margin, y + 10)
    c.setFont("Helvetica-Bold", 8)
    c.setFillColorRGB(0.145, 0.329, 0.871)
    c.drawString(margin, y + 14, "PaperFlow AI")
    c.setFillColorRGB(0.42, 0.45, 0.50)
    c.drawString(margin + 60, y + 14, f"  ·  {topic[:80]}")

    y -= 30

    # Title
    c.setFont("Helvetica-Bold", 16)
    c.setFillColorRGB(0.118, 0.227, 0.373)
    # Word-wrap title
    words = topic.split()
    line_buf = []
    for word in words:
        if c.stringWidth(" ".join(line_buf + [word]), "Helvetica-Bold", 16) < (w - 2 * margin):
            line_buf.append(word)
        else:
            c.drawString(margin, y, " ".join(line_buf))
            y -= 20
            line_buf = [word]
    if line_buf:
        c.drawString(margin, y, " ".join(line_buf))
        y -= 30

    # Meta
    c.setFont("Helvetica-Oblique", 8)
    c.setFillColorRGB(0.42, 0.45, 0.50)
    c.drawString(margin, y, f"PaperFlow AI · v{version}")
    y -= 22

    # Body
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            y -= 6
            continue
        if y < margin + 30:
            new_page()
            y = h - margin - 20

        if block.startswith("## "):
            c.setFont("Helvetica-Bold", 12)
            c.setFillColorRGB(0.118, 0.227, 0.373)
            c.drawString(margin, y, block[3:][:100])
            y -= 4
            c.setStrokeColorRGB(0.145, 0.329, 0.871)
            c.setLineWidth(0.8)
            c.line(margin, y, w - margin, y)
            y -= 14
        elif block.startswith("### "):
            c.setFont("Helvetica-Bold", 10)
            c.setFillColorRGB(0.118, 0.227, 0.373)
            c.drawString(margin, y, block[4:][:100])
            y -= 14
        else:
            c.setFont("Helvetica", 9.5)
            c.setFillColorRGB(0.10, 0.10, 0.12)
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", block)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            # Simple word-wrap
            words = clean.split()
            line_buf = []
            for word in words:
                test = " ".join(line_buf + [word])
                if c.stringWidth(test, "Helvetica", 9.5) < (w - 2 * margin):
                    line_buf.append(word)
                else:
                    c.drawString(margin, y, " ".join(line_buf))
                    y -= 13
                    if y < margin + 20:
                        new_page()
                        y = h - margin - 20
                    line_buf = [word]
            if line_buf:
                c.drawString(margin, y, " ".join(line_buf))
                y -= 13
            y -= 4

    c.save()
    logger.info(f"[Exporters] PDF reportlab fallback: {pdf_path}")
    return str(pdf_path.relative_to(storage_manager.base_dir))
