from __future__ import annotations

import csv
import io
import json
from pathlib import Path
from statistics import mean
from textwrap import wrap
from uuid import UUID

import httpx
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import settings
from app.core.storage import storage_manager
from app.models.analytics import AnalysisArtifact, AnalysisRun, Dataset, DatasetColumn, FigureArtifact


def _infer_dtype(series: pd.Series) -> str:
    import pandas as pd

    if pd.api.types.is_bool_dtype(series):
        return "boolean"
    if pd.api.types.is_integer_dtype(series):
        return "integer"
    if pd.api.types.is_numeric_dtype(series):
        return "number"
    return "string"


async def create_dataset(
    db: AsyncSession,
    *,
    project_id: UUID,
    title: str,
    description: str | None,
    rows: list[dict],
) -> Dataset:
    import pandas as pd

    if not rows:
        raise ValueError("Dataset rows cannot be empty")

    df = pd.DataFrame(rows)
    csv_bytes = df.to_csv(index=False).encode("utf-8")
    saved = await storage_manager.save_dataset_bytes(data=csv_bytes, filename=f"{title}.csv", project_id=project_id)

    dataset = Dataset(
        project_id=project_id,
        title=title,
        description=description,
        source_type="manual",
        file_path=saved["file_path"],
        row_count=len(df.index),
        column_count=len(df.columns),
        schema_json={"columns": list(df.columns)},
    )
    db.add(dataset)
    await db.flush()

    for position, column_name in enumerate(df.columns):
        series = df[column_name]
        summary = {
            "nulls": int(series.isna().sum()),
            "unique": int(series.nunique(dropna=True)),
        }
        if pd.api.types.is_numeric_dtype(series):
            numeric = pd.to_numeric(series, errors="coerce").dropna()
            if not numeric.empty:
                summary["mean"] = float(numeric.mean())
        db.add(
            DatasetColumn(
                dataset_id=dataset.id,
                name=str(column_name),
                data_type=_infer_dtype(series),
                position=position,
                is_nullable=bool(series.isna().any()),
                summary_json=summary,
            )
        )

    await db.commit()
    stmt = select(Dataset).where(Dataset.id == dataset.id).options(selectinload(Dataset.columns))
    hydrated = await db.execute(stmt)
    return hydrated.scalars().first()


def _load_dataset_frame(dataset: Dataset) -> pd.DataFrame:
    import pandas as pd

    if not dataset.file_path:
        return pd.DataFrame()
    with storage_manager.local_path(dataset.file_path, suffix=".csv") as local_path:
        return pd.read_csv(local_path)


def _local_analysis(df: pd.DataFrame, analysis_type: str, input_params: dict) -> dict:
    import pandas as pd

    analysis_type = analysis_type.lower()
    summary: dict = {"analysis_type": analysis_type, "row_count": int(len(df.index)), "column_count": int(len(df.columns))}
    warnings: list[str] = []

    if analysis_type == "descriptives":
        summary["columns"] = {
            column: {
                "type": _infer_dtype(df[column]),
                "non_null": int(df[column].notna().sum()),
                "unique": int(df[column].nunique(dropna=True)),
            }
            for column in df.columns
        }
    elif analysis_type == "group_comparison":
        group_col = input_params.get("group_column")
        value_col = input_params.get("value_column")
        if group_col in df.columns and value_col in df.columns:
            grouped = df.groupby(group_col)[value_col].agg(["count", "mean"]).reset_index()
            summary["groups"] = grouped.to_dict(orient="records")
        else:
            warnings.append("Missing group_column or value_column for group comparison")
    elif analysis_type in {"linear_regression", "logistic_regression"}:
        target = input_params.get("target_column")
        features = input_params.get("feature_columns") or []
        summary["target_column"] = target
        summary["feature_columns"] = features
        summary["note"] = "Regression fallback executed locally without coefficient fitting. Use r-engine for full model output."
        warnings.append("Local fallback does not fit coefficients; r-engine recommended.")
    else:
        summary["note"] = "Analysis type accepted but currently summarized via generic local fallback."
        warnings.append("Advanced analysis delegated to r-engine when available.")

    return {"summary": summary, "warnings": warnings, "script": f"# local fallback analysis\n# type: {analysis_type}\n# params: {json.dumps(input_params)}"}


def _analysis_payload(run: AnalysisRun) -> dict:
    return {
        "title": getattr(run, "title", None),
        "analysis_type": getattr(run, "analysis_type", None),
        "summary": getattr(run, "result_summary", None),
        "warnings": getattr(run, "warnings", None) or [],
        "script": getattr(run, "script_text", None),
    }


def _render_analysis_html(run: AnalysisRun) -> bytes:
    payload = _analysis_payload(run)
    body = (
        "<html><body><h1>{title}</h1><h2>{analysis_type}</h2><pre>{summary}</pre><pre>{script}</pre></body></html>".format(
            title=payload.get("title") or f"Analysis {run.id}",
            analysis_type=payload.get("analysis_type") or "analysis",
            summary=json.dumps(payload.get("summary") or {}, indent=2, ensure_ascii=False),
            script=payload.get("script") or "",
        )
    )
    return body.encode("utf-8")


def _render_analysis_docx(run: AnalysisRun) -> bytes:
    from docx import Document

    payload = _analysis_payload(run)
    doc = Document()
    doc.add_heading(payload.get("title") or f"Analysis {run.id}", level=1)
    doc.add_paragraph(f"Type: {payload.get('analysis_type') or 'analysis'}")
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(json.dumps(payload.get("summary") or {}, indent=2, ensure_ascii=False))
    warnings = payload.get("warnings") or []
    if warnings:
        doc.add_heading("Warnings", level=2)
        for warning in warnings:
            doc.add_paragraph(str(warning), style="List Bullet")
    if payload.get("script"):
        doc.add_heading("Script", level=2)
        doc.add_paragraph(str(payload["script"]))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _render_analysis_pdf(run: AnalysisRun) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas

    payload = _analysis_payload(run)
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    x = 2 * cm
    y = height - 2 * cm

    def write_block(text: str, *, font: str = "Helvetica", size: int = 10) -> None:
        nonlocal y
        pdf.setFont(font, size)
        for raw_line in (text or "").splitlines() or [""]:
            for line in wrap(raw_line, 100) or [""]:
                if y < 2 * cm:
                    pdf.showPage()
                    pdf.setFont(font, size)
                    y = height - 2 * cm
                pdf.drawString(x, y, line)
                y -= 0.55 * cm

    write_block(payload.get("title") or f"Analysis {run.id}", font="Helvetica-Bold", size=16)
    write_block(f"Type: {payload.get('analysis_type') or 'analysis'}", font="Helvetica", size=11)
    write_block("Summary", font="Helvetica-Bold", size=12)
    write_block(json.dumps(payload.get("summary") or {}, indent=2, ensure_ascii=False))
    warnings = payload.get("warnings") or []
    if warnings:
        write_block("Warnings", font="Helvetica-Bold", size=12)
        for warning in warnings:
            write_block(f"- {warning}")
    if payload.get("script"):
        write_block("Script", font="Helvetica-Bold", size=12)
        write_block(str(payload["script"]))

    pdf.save()
    return buffer.getvalue()


def _artifact_matches_format(artifact: AnalysisArtifact, fmt: str) -> bool:
    metadata = getattr(artifact, "metadata_json", None) or {}
    filename = str(getattr(artifact, "filename", "") or "").lower()
    return (
        getattr(artifact, "artifact_type", None) == f"report_{fmt}"
        or metadata.get("format") == fmt
        or filename.endswith(f".{fmt}")
    )


def _normalize_text_value(value: object | None) -> str | None:
    if value is None:
        return None
    if isinstance(value, str):
        return value
    if isinstance(value, list):
        return "\n\n".join(str(item) for item in value if item is not None)
    return str(value)


def _normalize_figure_payload(value: object | None) -> dict[str, object]:
    if not isinstance(value, dict):
        return {
            "title": "Summary figure placeholder",
            "caption": "Generated from analysis summary",
        }

    normalized = dict(value)
    normalized["title"] = _normalize_text_value(normalized.get("title")) or "Summary figure"
    normalized["caption"] = _normalize_text_value(normalized.get("caption"))
    if "analysis_type" in normalized:
        normalized["analysis_type"] = _normalize_text_value(normalized.get("analysis_type"))
    return normalized


async def _persist_export_artifact(
    db: AsyncSession,
    *,
    run: AnalysisRun,
    fmt: str,
    data: bytes,
    mime_type: str,
    artifact_manifest: list[dict],
) -> None:
    saved = await storage_manager.save_artifact_bytes(
        data=data,
        filename=f"{run.id}.{fmt}",
        project_id=run.project_id,
    )
    db.add(
        AnalysisArtifact(
            analysis_run_id=run.id,
            artifact_type=f"report_{fmt}",
            filename=saved["filename"],
            file_path=saved["file_path"],
            mime_type=mime_type,
            metadata_json={"analysis_type": run.analysis_type, "format": fmt},
        )
    )
    artifact_manifest.append(
        {
            "artifact_type": f"report_{fmt}",
            "filename": saved["filename"],
            "file_path": saved["file_path"],
        }
    )


async def create_analysis_run(
    db: AsyncSession,
    *,
    project_id: UUID,
    dataset: Dataset | None,
    title: str,
    analysis_type: str,
    input_params: dict,
) -> AnalysisRun:
    import pandas as pd

    rows: list[dict] = []
    if dataset is not None:
        df = _load_dataset_frame(dataset)
        rows = df.to_dict(orient="records")
    else:
        df = pd.DataFrame()

    payload = {
        "analysis_type": analysis_type,
        "input_params": input_params,
        "rows": rows,
    }
    response_data: dict | None = None
    warnings: list[str] = []
    try:
        async with httpx.AsyncClient(timeout=20) as client:
            resp = await client.post(f"{settings.R_ENGINE_URL}/run-analysis", json=payload)
            resp.raise_for_status()
            response_data = resp.json()
    except Exception:
        response_data = _local_analysis(df, analysis_type, input_params)
        warnings.append("r-engine unavailable, local fallback used")

    warnings.extend(response_data.get("warnings") or [])
    artifact_manifest: list[dict] = []
    run = AnalysisRun(
        project_id=project_id,
        dataset_id=dataset.id if dataset else None,
        title=title,
        analysis_type=analysis_type,
        status="completed",
        input_params=input_params,
        runtime_metadata={
            "engine": "r-engine" if "engine_version" in response_data else "local_fallback",
            "dataset_snapshot": {
                "dataset_id": str(dataset.id) if dataset else None,
                "row_count": len(rows),
                "column_count": len(rows[0].keys()) if rows else 0,
            },
            "artifact_manifest": artifact_manifest,
        },
        script_text=_normalize_text_value(response_data.get("script")),
        warnings=[str(item) for item in warnings],
        result_summary=response_data.get("summary"),
        engine_version=_normalize_text_value(response_data.get("engine_version")) or "local-fallback",
    )
    db.add(run)
    await db.flush()

    report_body = json.dumps(response_data.get("summary") or {}, indent=2, ensure_ascii=False)
    report_saved = await storage_manager.save_text_artifact(
        text=report_body,
        filename=f"{run.id}_summary.json",
        project_id=project_id,
    )
    db.add(
        AnalysisArtifact(
            analysis_run_id=run.id,
            artifact_type="summary",
            filename=report_saved["filename"],
            file_path=report_saved["file_path"],
            mime_type="application/json",
            metadata_json={"analysis_type": analysis_type},
        )
    )
    artifact_manifest.append({"artifact_type": "summary", "filename": report_saved["filename"], "file_path": report_saved["file_path"]})

    chart_data = _normalize_figure_payload(response_data.get("figure"))
    chart_saved = await storage_manager.save_text_artifact(
        text=json.dumps(chart_data, indent=2, ensure_ascii=False),
        filename=f"{run.id}_figure.json",
        project_id=project_id,
    )
    db.add(
        FigureArtifact(
            analysis_run_id=run.id,
            title=chart_data.get("title", "Summary figure"),
            caption=chart_data.get("caption"),
            filename=chart_saved["filename"],
            file_path=chart_saved["file_path"],
            metadata_json=chart_data,
        )
    )
    artifact_manifest.append({"artifact_type": "figure", "filename": chart_saved["filename"], "file_path": chart_saved["file_path"]})

    await _persist_export_artifact(
        db,
        run=run,
        fmt="html",
        data=_render_analysis_html(run),
        mime_type="text/html",
        artifact_manifest=artifact_manifest,
    )
    await _persist_export_artifact(
        db,
        run=run,
        fmt="pdf",
        data=_render_analysis_pdf(run),
        mime_type="application/pdf",
        artifact_manifest=artifact_manifest,
    )
    await _persist_export_artifact(
        db,
        run=run,
        fmt="docx",
        data=_render_analysis_docx(run),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        artifact_manifest=artifact_manifest,
    )

    await db.commit()
    stmt = (
        select(AnalysisRun)
        .where(AnalysisRun.id == run.id)
        .options(selectinload(AnalysisRun.artifacts))
    )
    hydrated = await db.execute(stmt)
    return hydrated.scalars().first()


async def export_analysis_run(db: AsyncSession, *, run: AnalysisRun, fmt: str) -> tuple[bytes, str, str]:
    if fmt not in {"html", "pdf", "docx"}:
        raise ValueError("Unsupported export format")

    if getattr(run, "status", None) != "completed":
        raise ValueError(f"Analysis run {getattr(run, 'id', 'unknown')} is not completed")

    artifacts = list(getattr(run, "artifacts", None) or [])
    artifact = next((item for item in artifacts if _artifact_matches_format(item, fmt)), None)
    if artifact is None:
        raise ValueError(f"No persisted {fmt} artifact found for analysis run {getattr(run, 'id', 'unknown')}")

    try:
        data = storage_manager.read_bytes(artifact.file_path)
    except FileNotFoundError as exc:
        raise ValueError(f"Artifact {artifact.file_path} missing from storage") from exc

    return (
        data,
        getattr(artifact, "mime_type", None)
        or {
            "html": "text/html",
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }[fmt],
        getattr(artifact, "filename", None) or f"{run.id}.{fmt}",
    )
