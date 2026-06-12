"""Writing document export utilities (Markdown, DOCX, PDF).

Reads a `WritingDocument` (with sections and claim_links eager-loaded) and emits
binary content suitable for download, plus a final auto-generated bibliography
in the requested citation style (apa | vancouver | ama).
"""

from __future__ import annotations

import io
import re
from typing import Any, Iterable

from app.models.reference_item import ReferenceItem
from app.models.writing import WritingDocument


CITATION_STYLES = {"apa", "vancouver", "ama"}


# --------------------------------------------------------------------------- #
# Bibliography formatting
# --------------------------------------------------------------------------- #


def _format_authors_apa(authors: list[str] | None) -> str:
    if not authors:
        return ""
    return ", ".join(authors)


def _format_authors_vancouver(authors: list[str] | None) -> str:
    if not authors:
        return ""
    if len(authors) <= 6:
        return ", ".join(authors) + "."
    return ", ".join(authors[:6]) + ", et al."


def format_reference(item: ReferenceItem, style: str = "apa") -> str:
    style = style.lower() if style else "apa"
    if style not in CITATION_STYLES:
        style = "apa"
    title = (item.title or "Untitled").strip().rstrip(".")
    journal = (item.journal or "").strip()
    year = item.publication_year or "n.d."
    doi = item.doi
    pmid = item.pmid
    if style == "apa":
        authors = _format_authors_apa(item.authors)
        parts = []
        if authors:
            parts.append(authors)
        parts.append(f"({year}).")
        parts.append(f"{title}.")
        if journal:
            parts.append(f"*{journal}*.")
        if doi:
            parts.append(f"https://doi.org/{doi}")
        return " ".join(parts)
    if style == "vancouver":
        authors = _format_authors_vancouver(item.authors)
        parts = []
        if authors:
            parts.append(authors)
        parts.append(f"{title}.")
        if journal:
            parts.append(f"{journal}.")
        parts.append(f"{year}.")
        if doi:
            parts.append(f"doi:{doi}")
        elif pmid:
            parts.append(f"PMID: {pmid}")
        return " ".join(parts)
    # AMA
    authors = _format_authors_vancouver(item.authors)
    parts = []
    if authors:
        parts.append(authors)
    parts.append(f"{title}.")
    if journal:
        parts.append(f"*{journal}*.")
    parts.append(f"{year}.")
    if doi:
        parts.append(f"doi:{doi}")
    return " ".join(parts)


def build_bibliography_section(
    references: Iterable[ReferenceItem],
    *,
    style: str = "apa",
) -> str:
    """Return a Markdown bibliography block."""
    refs = list(references)
    if not refs:
        return ""
    lines = ["## References", ""]
    for idx, ref in enumerate(refs, start=1):
        lines.append(f"{idx}. {format_reference(ref, style=style)}")
    lines.append("")
    return "\n".join(lines)


# --------------------------------------------------------------------------- #
# Markdown
# --------------------------------------------------------------------------- #


def build_markdown(
    document: WritingDocument,
    *,
    references: Iterable[ReferenceItem] | None = None,
    citation_style: str = "apa",
) -> str:
    """Concatenate sections and append a bibliography in the chosen style."""
    parts: list[str] = []
    parts.append(f"# {document.title}\n")
    if document.mode:
        parts.append(f"*Mode: {document.mode} \u00b7 v{document.version}*\n")
    sections_sorted = sorted(
        document.sections,
        key=lambda item: (item.position or 0, item.created_at),
    )
    for section in sections_sorted:
        body = (section.content_markdown or "").strip()
        if not body:
            parts.append(f"## {section.heading}\n\n_(empty)_\n")
        else:
            # If the body already starts with a "## " heading we let it be,
            # otherwise prepend the heading so output is self-contained.
            if not body.startswith("#"):
                parts.append(f"## {section.heading}\n\n{body}\n")
            else:
                parts.append(body + "\n")
    if references is not None:
        bib = build_bibliography_section(references, style=citation_style)
        if bib:
            parts.append(bib)
    return "\n".join(parts).strip() + "\n"


# --------------------------------------------------------------------------- #
# DOCX (python-docx)
# --------------------------------------------------------------------------- #

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_LIST_RE = re.compile(r"^[\-\*]\s+(.*)$")
_NUM_LIST_RE = re.compile(r"^\d+\.\s+(.*)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")
_CITATION_RE = re.compile(r"(\[[A-Z]\d+\])")

_INLINE_PATTERN_RE = re.compile(
    r"(\*\*(?P<bold>.+?)\*\*)"
    r"|(?<!\*)\*(?P<italic>[^*\n]+)\*(?!\*)"
    r"|(?P<cite>\[[A-Z]\d+\])"
)


def _add_inline_runs(paragraph, text: str) -> None:
    """Add runs to a docx paragraph honoring **bold**, *italic* and [A1] citations."""
    # Tokenize into bold / italic / citation / plain spans.
    cursor = 0
    for match in _INLINE_PATTERN_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        if match.group("bold"):
            run = paragraph.add_run(match.group("bold"))
            run.bold = True
        elif match.group("italic"):
            run = paragraph.add_run(match.group("italic"))
            run.italic = True
        elif match.group("cite"):
            run = paragraph.add_run(match.group("cite"))
            run.font.superscript = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def build_docx(
    document: WritingDocument,
    *,
    references: Iterable[ReferenceItem] | None = None,
    citation_style: str = "apa",
) -> bytes:
    """Render the document as a .docx binary blob."""
    from docx import Document  # type: ignore

    doc = Document()
    doc.add_heading(document.title, level=0)
    if document.mode:
        meta = doc.add_paragraph()
        run = meta.add_run(f"Mode: {document.mode} \u00b7 v{document.version}")
        run.italic = True

    sections_sorted = sorted(
        document.sections,
        key=lambda item: (item.position or 0, item.created_at),
    )

    for section in sections_sorted:
        doc.add_heading(section.heading, level=1)
        body = (section.content_markdown or "").strip()
        if not body:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("(empty)")
            run.italic = True
            continue

        for raw_line in body.splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                doc.add_paragraph("")
                continue

            heading_match = _HEADING_RE.match(line)
            if heading_match:
                hashes, heading_text = heading_match.groups()
                level = min(len(hashes) + 1, 4)  # leave 0/1 for title and section
                # Avoid duplicating the section heading if the LLM already included it.
                if (
                    heading_text.strip().lower() == section.heading.strip().lower()
                    and level <= 2
                ):
                    continue
                doc.add_heading(heading_text.strip(), level=level)
                continue

            list_match = _LIST_RE.match(line)
            if list_match:
                paragraph = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(paragraph, list_match.group(1))
                continue

            num_match = _NUM_LIST_RE.match(line)
            if num_match:
                paragraph = doc.add_paragraph(style="List Number")
                _add_inline_runs(paragraph, num_match.group(1))
                continue

            paragraph = doc.add_paragraph()
            _add_inline_runs(paragraph, line)

    refs = list(references) if references is not None else []
    if refs:
        doc.add_heading("References", level=1)
        for idx, ref in enumerate(refs, start=1):
            paragraph = doc.add_paragraph(style="List Number")
            _add_inline_runs(paragraph, format_reference(ref, style=citation_style))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
# PDF (reportlab)
# --------------------------------------------------------------------------- #


def build_pdf(
    document: WritingDocument,
    *,
    references: Iterable[ReferenceItem] | None = None,
    citation_style: str = "apa",
) -> bytes:
    """Render the document as a PDF binary blob using reportlab."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        ListFlowable,
        ListItem,
    )

    buffer = io.BytesIO()
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        rightMargin=2.54 * cm,
        leftMargin=2.54 * cm,
        topMargin=2.54 * cm,
        bottomMargin=2.54 * cm,
        title=document.title,
    )
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="PFTitle",
            parent=styles["Title"],
            fontName="Times-Bold",
            fontSize=20,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="PFHeading",
            parent=styles["Heading2"],
            fontName="Times-Bold",
            fontSize=14,
            spaceBefore=14,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="PFSubHeading",
            parent=styles["Heading3"],
            fontName="Times-Bold",
            fontSize=12,
            spaceBefore=10,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="PFBody",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=11,
            leading=15,
            spaceAfter=6,
        )
    )

    def md_inline_to_html(text: str) -> str:
        # Simple markdown-to-mini-html for reportlab Paragraph.
        out = text
        out = _BOLD_RE.sub(r"<b>\1</b>", out)
        out = _ITALIC_RE.sub(r"<i>\1</i>", out)
        out = _CITATION_RE.sub(r"<super>\1</super>", out)
        return out

    story: list[Any] = []
    story.append(Paragraph(document.title, styles["PFTitle"]))
    if document.mode:
        story.append(
            Paragraph(
                f"<i>Mode: {document.mode} \u00b7 v{document.version}</i>",
                styles["PFBody"],
            )
        )
        story.append(Spacer(1, 8))

    sections_sorted = sorted(
        document.sections,
        key=lambda item: (item.position or 0, item.created_at),
    )
    for section in sections_sorted:
        story.append(Paragraph(section.heading, styles["PFHeading"]))
        body = (section.content_markdown or "").strip()
        if not body:
            story.append(Paragraph("<i>(empty)</i>", styles["PFBody"]))
            continue

        bullet_buffer: list[str] = []
        number_buffer: list[str] = []

        def flush_bullets() -> None:
            if bullet_buffer:
                story.append(
                    ListFlowable(
                        [
                            ListItem(
                                Paragraph(md_inline_to_html(item), styles["PFBody"])
                            )
                            for item in bullet_buffer
                        ],
                        bulletType="bullet",
                        leftIndent=14,
                    )
                )
                bullet_buffer.clear()

        def flush_numbers() -> None:
            if number_buffer:
                story.append(
                    ListFlowable(
                        [
                            ListItem(
                                Paragraph(md_inline_to_html(item), styles["PFBody"])
                            )
                            for item in number_buffer
                        ],
                        bulletType="1",
                        leftIndent=14,
                    )
                )
                number_buffer.clear()

        for raw_line in body.splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                flush_bullets()
                flush_numbers()
                story.append(Spacer(1, 4))
                continue

            heading_match = _HEADING_RE.match(line)
            if heading_match:
                flush_bullets()
                flush_numbers()
                hashes, heading_text = heading_match.groups()
                if (
                    heading_text.strip().lower() == section.heading.strip().lower()
                    and len(hashes) <= 2
                ):
                    continue
                story.append(
                    Paragraph(
                        md_inline_to_html(heading_text.strip()), styles["PFSubHeading"]
                    )
                )
                continue

            list_match = _LIST_RE.match(line)
            if list_match:
                flush_numbers()
                bullet_buffer.append(list_match.group(1))
                continue

            num_match = _NUM_LIST_RE.match(line)
            if num_match:
                flush_bullets()
                number_buffer.append(num_match.group(1))
                continue

            flush_bullets()
            flush_numbers()
            story.append(Paragraph(md_inline_to_html(line), styles["PFBody"]))

        flush_bullets()
        flush_numbers()

    refs = list(references) if references is not None else []
    if refs:
        story.append(Paragraph("References", styles["PFHeading"]))
        story.append(
            ListFlowable(
                [
                    ListItem(
                        Paragraph(
                            md_inline_to_html(
                                format_reference(ref, style=citation_style)
                            ),
                            styles["PFBody"],
                        )
                    )
                    for ref in refs
                ],
                bulletType="1",
                leftIndent=14,
            )
        )

    pdf.build(story)
    return buffer.getvalue()
